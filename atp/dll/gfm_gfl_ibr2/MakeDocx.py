# Copyright (C) 2025 Meltran, Inc
# file: MakeDocx.py
""" Assembles PNG plots into a document.

After the script runs, open the DocX file in Word, select
all, then Ctrl-F9 to update fields. The figure numbers, initially
blank, should then appear.

Public Functions:
    :main: does the work
"""

#import glob
#import csv
import json
from docx import Document
from docx.shared import Inches
from docx.enum.section import WD_ORIENT

bWide = False

from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def MarkIndexEntry(entry,paragraph):
  run = paragraph.add_run()
  r = run._r
  fldChar = OxmlElement('w:fldChar')
  fldChar.set(qn('w:fldCharType'), 'begin')
  r.append(fldChar)

  run = paragraph.add_run()
  r = run._r
  instrText = OxmlElement('w:instrText')
  instrText.set(qn('xml:space'), 'preserve')
  instrText.text = ' XE "%s" '%(entry)
  r.append(instrText)

  run = paragraph.add_run()
  r = run._r
  fldChar = OxmlElement('w:fldChar')
  fldChar.set(qn('w:fldCharType'), 'end')
  r.append(fldChar)

def Figure(paragraph):
  run = paragraph.add_run()
  r = run._r
  fldChar = OxmlElement('w:fldChar')
  fldChar.set(qn('w:fldCharType'), 'begin')
  r.append(fldChar)
  instrText = OxmlElement('w:instrText')
  instrText.text = ' SEQ Figure \* ARABIC'
  r.append(instrText)
  fldChar = OxmlElement('w:fldChar')
  fldChar.set(qn('w:fldCharType'), 'end')
  r.append(fldChar)

def Table(paragraph):
  run = run = paragraph.add_run()
  r = run._r
  fldChar = OxmlElement('w:fldChar')
  fldChar.set(qn('w:fldCharType'), 'begin')
  r.append(fldChar)
  instrText = OxmlElement('w:instrText')
  instrText.text = ' SEQ Table \* ARABIC'
  r.append(instrText)
  fldChar = OxmlElement('w:fldChar')
  fldChar.set(qn('w:fldCharType'), 'end')
  r.append(fldChar)

def cellval(tok):
  if float(tok) > 0:
    return tok
  return 'n/a'

document = Document()
if bWide:
  for section in document.sections:
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
else:
  for section in document.sections:
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)

#table = document.add_table(rows=1, cols=8)
#hdr_cells = table.rows[0].cells
#hdr_cells[0].text = 'Case'
#hdr_cells[1].text = 'Site'
#hdr_cells[2].text = 'Tfault'
#hdr_cells[3].text = 'TD32F'
#hdr_cells[4].text = 'OC21P'
#hdr_cells[5].text = 'OC21G'
#hdr_cells[6].text = 'TD21P'
#hdr_cells[7].text = 'TD21G'
#with open('T400LTimes.csv', mode='r') as infile:
#    reader = csv.reader(infile)
#    for row in reader:
#        png_name = row[0]
#        site_name = row[1]
#        tfault = row[2]
#        row_cells = table.add_row().cells
#        row_cells[0].text = png_name
#        row_cells[1].text = site_name
#        row_cells[2].text = tfault
#        for i in range (3, 8):
#            row_cells[i].text = cellval(row[i])
#document.add_page_break()

#files = sorted(glob.glob ('*.png'))
with open('cases.json', 'r') as f:
  cases = json.load(f)

bPageBreak = False # page break after every two plots

for idx in range(len(cases)):
  fignum = cases[idx]['Fig']
  title = cases[idx]['Title']
  fname = 'fig{:d}.png'.format(fignum)
  if bWide:
    document.add_picture(fname, width=Inches(9.5))
  else:
    document.add_picture(fname, width=Inches(6.5))
  # adding a cross-reference enabled caption
  paragraph = document.add_paragraph('Figure ', style='Caption')
  Figure (paragraph)
  paragraph.add_run(': {:s}'.format (title))
  #document.add_paragraph('Figure ' + str(fignum) + ': ' + title, style=document.styles['Caption'])
  if bPageBreak:
#    document.add_page_break()
    bPageBreak = False
  else:
    bPageBreak = True

document.save('plots_new.docx')
